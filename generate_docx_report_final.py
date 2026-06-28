"""
AUCA Big Data Analytics Final Project
Professional Technical Report (DOCX)
Author: BYIRINGIRO Elie Yvan
Registration Number: 101219
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# SETUP
doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 61, 130)
    return p

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    if level == 1:
        h_format = h.runs[0].font
        h_format.size = Pt(16)
        h_format.bold = True
        h_format.color.rgb = RGBColor(0, 95, 163)

def add_paragraph(doc, text, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    return p

# TITLE PAGE
add_title(doc, "AUCA Big Data Analytics")
doc.add_paragraph()
add_title(doc, "Final Project Report")
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph("Customer Behavior Prediction & E-Commerce Analytics Pipeline")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.runs[0]
run.font.size = Pt(14)
run.font.italic = True
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.add_run("Author: ").bold = True
info.add_run("BYIRINGIRO Elie Yvan\n")
info.add_run("Registration Number: ").bold = True
info.add_run("101219\n")
info.add_run("Submission Date: ").bold = True
info.add_run("June 21, 2026\n")
info.add_run("Report Generated: ").bold = True
info.add_run(datetime.now().strftime("%B %d, %Y\n"))
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
tech_stack = doc.add_paragraph()
tech_stack.add_run("Technology Stack: ").bold = True
tech_stack.add_run("MongoDB • HBase • Apache Spark (PySpark) • Python 3.14 • Ubuntu WSL2")
tech_stack.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# TABLE OF CONTENTS
add_heading(doc, "Table of Contents")
toc_items = ["1. Executive Summary", "2. Introduction & Objectives", "3. Architecture & Technology Stack", 
             "4. Data Modeling & Design", "5. Implementation Summary", "6. Results & Findings",
             "7. Visualizations & Insights", "8. Customer Segmentation Analysis", 
             "9. Technical Challenges & Solutions", "10. Conclusion & Future Work"]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# EXECUTIVE SUMMARY
add_heading(doc, "1. Executive Summary")
add_paragraph(doc, 
    "This project demonstrates an end-to-end big data analytics pipeline integrating MongoDB, HBase, and Apache Spark "
    "to analyze e-commerce customer behavior. The pipeline processed 2,000 transactions across 500 products for 500 users, "
    "generating actionable insights through Customer Lifetime Value (CLV) analysis.")

add_paragraph(doc, "Key Outcomes:")
for outcome in ["Designed scalable data models across multiple NoSQL and distributed systems",
                "Loaded and aggregated 2,260+ records across MongoDB and HBase",
                "Executed 4 advanced Spark SQL queries on normalized transaction data",
                "Computed CLV for 327 customers with 4-tier segmentation strategy",
                "Generated 4 publication-ready visualizations from real analytics"]:
    doc.add_paragraph(outcome, style='List Bullet')

doc.add_page_break()

# INTRODUCTION
add_heading(doc, "2. Introduction & Objectives")
add_paragraph(doc,
    "The modern e-commerce landscape demands sophisticated data engineering capabilities. This project integrates "
    "disparate data systems (MongoDB, HBase, Spark) into a cohesive analytics platform.")

add_heading(doc, "2.1 Project Objectives", level=2)
for obj in ["Design and implement data models for MongoDB, HBase, and Spark",
            "Demonstrate distributed batch processing and SQL query execution",
            "Perform integrated analytics combining multiple data sources",
            "Segment customers by lifetime value and purchasing behavior",
            "Visualize key business metrics for decision-making"]:
    doc.add_paragraph(obj, style='List Bullet')

add_heading(doc, "2.2 Dataset", level=2)
dataset = [["Dimension", "Count"], ["Users", "500"], ["Categories", "20"], ["Products", "500"], 
           ["Sessions", "3,000"], ["Transactions", "2,000"]]
table = doc.add_table(rows=len(dataset), cols=2)
table.style = 'Light Grid Accent 1'
for i, row_data in enumerate(dataset):
    table.rows[i].cells[0].text = row_data[0]
    table.rows[i].cells[1].text = row_data[1]
    if i == 0:
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[i].cells[1].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# ARCHITECTURE
add_heading(doc, "3. Architecture & Technology Stack")
add_heading(doc, "3.1 MongoDB", level=2)
add_paragraph(doc, "Document-oriented NoSQL database storing 2,260 records across 4 collections (users, categories, products, transactions). "
              "Aggregation pipelines enable complex analytical queries without ETL.")

add_heading(doc, "3.2 HBase", level=2)
add_paragraph(doc, "Wide-column store simulation for time-series session data and product metrics. Row-key design enables efficient range queries. "
              "Two tables: user_sessions (3,000 rows), product_performance (9,428 rows).")

add_heading(doc, "3.3 Apache Spark", level=2)
add_paragraph(doc, "Distributed computing framework handling data normalization, aggregation, and joins. PySpark processes 2,000 transactions in <30 seconds. "
              "Version 4.1.2 with Python 3.14.")

tech_specs = [["Component", "Technology"], ["Database 1", "MongoDB 8.0.26"], ["Database 2", "HBase (Simulation)"],
              ["Processing", "Apache Spark 4.1.2"], ["Language", "Python 3.14"], ["Visualization", "Matplotlib 3.11 + Seaborn 0.13"]]
table = doc.add_table(rows=len(tech_specs), cols=2)
table.style = 'Light Grid Accent 1'
for i, row_data in enumerate(tech_specs):
    table.rows[i].cells[0].text = row_data[0]
    table.rows[i].cells[1].text = row_data[1]
    if i == 0:
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[i].cells[1].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# RESULTS
add_heading(doc, "4. Results & Findings")

add_heading(doc, "4.1 Customer Value Distribution", level=2)
segmentation = [["Segment", "Count", "% of Customers", "Avg CLV", "Total Revenue"],
                ["Premium", "29", "8.9%", "$2,502", "$100,606"],
                ["High Value", "111", "33.9%", "$1,518", "$278,497"],
                ["Medium Value", "122", "37.3%", "$903", "$198,762"],
                ["Low Value", "64", "19.6%", "$338", "$29,474"]]
table = doc.add_table(rows=len(segmentation), cols=5)
table.style = 'Light Grid Accent 1'
for i, row_data in enumerate(segmentation):
    for j, cell in enumerate(row_data):
        table.rows[i].cells[j].text = cell
        if i == 0:
            table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

add_paragraph(doc, "\nKey Insight: Top 8.9% of customers generate 36.1% of revenue (Pareto principle). These are high-priority for retention.")

add_heading(doc, "4.2 Category & Payment Analysis", level=2)
add_paragraph(doc, "Revenue Distribution:\n"
              "  • cat_004 leads at $43,956 | cat_013 lowest at $16,653\n"
              "  • Insight: Order volume drives revenue, not price per unit\n\n"
              "Payment Methods: All 5 methods evenly adopted (~20% each). Google Pay leads revenue at $131.7K.")

add_heading(doc, "4.3 Device Behavior", level=2)
add_paragraph(doc, "  • Tablet: 1,036 sessions (34.5%) | 933 sec avg\n"
              "  • Desktop: 995 sessions (33.2%) | 948 sec avg\n"
              "  • Mobile: 969 sessions (32.3%) | 954 sec avg\n"
              "  • Insight: Tablet slightly preferred; engagement consistent across devices")

doc.add_page_break()

# VISUALIZATIONS
add_heading(doc, "5. Visualizations")

for plot_file, title in [("plot_1_clv_distribution.png", "Figure 1: CLV Distribution"),
                         ("plot_2_revenue_by_category.png", "Figure 2: Revenue by Category"),
                         ("plot_3_payment_methods.png", "Figure 3: Payment Methods"),
                         ("plot_4_device_usage.png", "Figure 4: Device Usage")]:
    add_heading(doc, title, level=2)
    try:
        doc.add_picture(plot_file, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        add_paragraph(doc, f"[{title} visualization available in project folder]")

doc.add_page_break()

# TECHNICAL CHALLENGES
add_heading(doc, "6. Technical Challenges & Solutions")

challenges = [
    ("WSL2 Package Management",
     "Python blocked package installation with 'externally-managed-environment' error. Used --break-system-packages flag which safely overrides protection in development environments."),
    
    ("PySpark Installation & /tmp Space",
     "455MB PySpark download timed out, then wheel build failed with 'No space left on device.' Increased /tmp from 1.9GB to 4GB with: sudo mount -o remount,size=4G /tmp. Reinstalled with success."),
    
    ("MongoDB ObjectId Serialization",
     "json.dumps() failed on MongoDB's ObjectId field when loading to Spark. Stripped _id fields before serialization: for t in transactions: if '_id' in t: del t['_id']. Solution preserved all analytical data while making JSON compatible."),
    
    ("Spark JSON Parsing (Windows Files)",
     "Spark 4.x read transaction.json as corrupt records due to Windows line endings. Added multiLine='true' option to Spark's reader, which handles pretty-printed JSON and line-ending inconsistencies."),
    
    ("HBase Unavailable",
     "No Docker or Hadoop cluster available. Built HBase as Python simulation preserving wide-column architecture, row-key design, and range queries. Project spec allows this; demonstrates architectural understanding without infrastructure."),
    
    ("matplotlib/seaborn Import Errors",
     "After install, modules still not found. Reinstalled to correct Python path using --break-system-packages. Pip's caching mechanism used cached wheels on second attempt."),
    
    ("WSL2↔Windows File Sync Delays",
     "File created in Notepad on Windows wasn't visible in WSL2 terminal (sync delay). Recreated directly in WSL2 via nano to avoid filesystem layer, ensuring consistent encoding and line endings."),
    
    ("Network Timeouts on Large Downloads",
     "matplotlib (11.1MB) and pillow (7.1MB) downloads timed out repeatedly. Pip's cache and retry mechanism meant subsequent runs resumed. Eventually all bytes transferred successfully.")
]

for i, (title, description) in enumerate(challenges, 1):
    p = doc.add_paragraph()
    p.add_run(f"Challenge {i}: {title}\n").bold = True
    p.add_run(f"What Happened & Solution: {description}")
    p.paragraph_format.space_after = Pt(12)

doc.add_page_break()

# CONCLUSION
add_heading(doc, "7. Conclusion")
add_paragraph(doc, "This project successfully integrated MongoDB, HBase, and Spark into a production-grade analytics pipeline. "
              "CLV analysis revealed actionable customer segmentation: 29 Premium customers drive 36% of revenue. "
              "Technical challenges were solved systematically through research and persistence.")

add_heading(doc, "7.1 Key Achievements", level=2)
for achievement in ["Designed scalable schemas across 3 data systems",
                    "Processed 5,760+ records in <30 seconds",
                    "Computed CLV for 327 customers with 4-tier segmentation",
                    "Generated 4 publication-ready visualizations",
                    "Overcame 8 real-world technical challenges"]:
    doc.add_paragraph(achievement, style='List Bullet')

add_paragraph(doc, "\n" + "="*70)
add_paragraph(doc, f"Completed by: BYIRINGIRO Elie Yvan (Reg. 101219)")
add_paragraph(doc, f"Date: {datetime.now().strftime('%B %d, %Y')}")
add_paragraph(doc, "="*70)

# SAVE
output_file = "AUCA_BigData_Final_Project_Report.docx"
doc.save(output_file)

print("=" * 70)
print("✓ Professional Report Generated!")
print("=" * 70)
print(f"File: {output_file}")
print(f"Author: BYIRINGIRO Elie Yvan (Reg. 101219)")
print(f"Pages: ~12")
print(f"Sections: Executive Summary, Architecture, Results, Visualizations, 8 Challenges, Conclusion")
print("=" * 70)